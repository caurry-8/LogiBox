from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from utils.data_store import DataStore


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def _add_key_value_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "指标"
    table.rows[0].cells[1].text = "结果"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = str(value)


def generate_word_report(filename: str, store: DataStore) -> str:
    if not store.has_data():
        raise ValueError("当前没有可用于生成报告的数据。")

    path = Path(filename)
    if path.suffix.lower() != ".docx":
        path = path.with_suffix(".docx")

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LogiBox V3.2 物流数据分析报告")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"数据源：{store.filename_only()}")

    _add_heading(document, "一、数据概况")
    _add_key_value_table(
        document,
        [
            ("数据行数", str(store.rows())),
            ("字段数量", str(store.cols())),
            ("缺失单元格", str(store.missing_cells())),
            ("重复行", str(store.duplicate_rows())),
        ],
    )

    _add_heading(document, "二、ABC 分类分析")
    abc = store.get_analysis("abc")
    if abc:
        counts = abc.get("counts", {})
        contributions = abc.get("contributions", {})
        _add_key_value_table(
            document,
            [
                ("分析字段", abc.get("value_column", "")),
                ("A 类 SKU", counts.get("A", 0)),
                ("B 类 SKU", counts.get("B", 0)),
                ("C 类 SKU", counts.get("C", 0)),
                ("A 类金额贡献", f"{contributions.get('A', 0):.2%}"),
                ("B 类金额贡献", f"{contributions.get('B', 0):.2%}"),
                ("C 类金额贡献", f"{contributions.get('C', 0):.2%}"),
            ],
        )
    else:
        document.add_paragraph("当前尚未完成 ABC 分类分析。")

    _add_heading(document, "三、XYZ 稳定性分析")
    xyz = store.get_analysis("xyz")
    if xyz:
        counts = xyz.get("counts", {})
        _add_key_value_table(
            document,
            [
                ("历史周期字段", ", ".join(xyz.get("period_columns", []))),
                ("X 类 SKU", counts.get("X", 0)),
                ("Y 类 SKU", counts.get("Y", 0)),
                ("Z 类 SKU", counts.get("Z", 0)),
                ("平均 CV", f"{xyz.get('mean_cv', 0):.4f}"),
            ],
        )
    else:
        document.add_paragraph("当前尚未完成 XYZ 稳定性分析。")

    _add_heading(document, "四、EOQ 计算结果")
    eoq = store.get_analysis("eoq")
    if eoq:
        rows = [(key, value) for key, value in eoq.items()]
        _add_key_value_table(document, rows)
    else:
        document.add_paragraph("当前尚未完成 EOQ 计算。")

    _add_heading(document, "五、安全库存与再订货点")
    safety = store.get_analysis("safety")
    if safety:
        _add_key_value_table(
            document,
            [
                ("安全库存", safety.get("安全库存", "")),
                ("再订货点 ROP", safety.get("再订货点 ROP", "")),
                ("Z 值", safety.get("Z 值", "")),
            ],
        )
    else:
        document.add_paragraph("当前尚未完成安全库存计算。")

    _add_heading(document, "六、分析建议")
    recommendations = [
        "优先关注 ABC 中 A 类库存的缺货风险与补货策略。",
        "对 XYZ 中 Z 类库存强化需求预测和安全库存复核。",
        "结合 EOQ 与安全库存结果优化订货批量和再订货点。",
        "正式决策前建议结合业务周期、供应商交期和仓储容量进行校验。",
    ]
    for item in recommendations:
        document.add_paragraph(item, style="List Bullet")

    document.save(path)
    return str(path)
